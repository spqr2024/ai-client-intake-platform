"""Document ingestion: text extraction and chunking.

Supported formats: PDF, DOCX, Markdown, plain text. PDF/DOCX extraction uses
optional dependencies (`pypdf`, `python-docx`); when they are absent the
uploader returns a clear, actionable error instead of crashing, and the rest
of the KB keeps working. Markdown and text need no dependencies at all.

Chunking is paragraph-aware with a character budget and overlap: retrieval
quality on long documents depends far more on sensible passage boundaries
than on the embedding model.
"""

import io
import logging
import re

logger = logging.getLogger(__name__)

CHUNK_TARGET_CHARS = 900
CHUNK_OVERLAP_CHARS = 150
MIN_CHUNK_CHARS = 80

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".md": "md",
    ".markdown": "md",
    ".txt": "txt",
    ".text": "txt",
}


class UnsupportedDocument(Exception):
    """Raised when a file type is unknown or its extractor is unavailable."""


def _extract_pdf(data: bytes) -> tuple[str, dict]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - depends on environment
        raise UnsupportedDocument(
            "PDF support requires the 'pypdf' package. Install it with "
            "`pip install pypdf`, or paste the text manually."
        ) from exc
    reader = PdfReader(io.BytesIO(data))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(p for p in pages if p)
    if not text.strip():
        raise UnsupportedDocument(
            "No extractable text found — this PDF is likely a scanned image. "
            "OCR is not performed; please paste the text manually."
        )
    return text, {"pages": len(reader.pages)}


def _extract_docx(data: bytes) -> tuple[str, dict]:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - depends on environment
        raise UnsupportedDocument(
            "DOCX support requires the 'python-docx' package. Install it with "
            "`pip install python-docx`, or paste the text manually."
        ) from exc
    document = docx.Document(io.BytesIO(data))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    text = "\n\n".join(paragraphs)
    if not text.strip():
        raise UnsupportedDocument("The document appears to be empty.")
    return text, {"paragraphs": len(paragraphs)}


def _strip_markdown(text: str) -> str:
    text = re.sub(r"```.*?```", " ", text, flags=re.DOTALL)  # code fences
    text = re.sub(r"!\[[^\]]*\]\([^)]*\)", " ", text)  # images
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)  # links → label
    text = re.sub(r"^\s{0,3}#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"[*_`>]", "", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def extract(filename: str, data: bytes) -> tuple[str, str, dict]:
    """Return (source_type, text, metadata) for an uploaded document."""
    suffix = ""
    if "." in filename:
        suffix = filename[filename.rfind(".") :].lower()
    source_type = SUPPORTED_EXTENSIONS.get(suffix)
    if source_type is None:
        raise UnsupportedDocument(
            f"Unsupported file type '{suffix or filename}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if source_type == "pdf":
        text, metadata = _extract_pdf(data)
    elif source_type == "docx":
        text, metadata = _extract_docx(data)
    else:
        try:
            raw = data.decode("utf-8")
        except UnicodeDecodeError:
            raw = data.decode("latin-1", errors="replace")
        text = _strip_markdown(raw) if source_type == "md" else raw.strip()
        metadata = {}
        if not text:
            raise UnsupportedDocument("The file is empty.")

    metadata.update({"bytes": len(data), "characters": len(text)})
    return source_type, text, metadata


def chunk_text(text: str, target: int = CHUNK_TARGET_CHARS, overlap: int = CHUNK_OVERLAP_CHARS) -> list[str]:
    """Split text into overlapping, paragraph-aligned passages."""
    text = (text or "").strip()
    if not text:
        return []
    if len(text) <= target:
        return [text]

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[str] = []
    current = ""

    for paragraph in paragraphs:
        # A single oversized paragraph is split on sentence boundaries.
        if len(paragraph) > target:
            if current:
                chunks.append(current)
                current = ""
            sentences = re.split(r"(?<=[.!?])\s+", paragraph)
            buffer = ""
            for sentence in sentences:
                if buffer and len(buffer) + len(sentence) + 1 > target:
                    chunks.append(buffer.strip())
                    buffer = buffer[-overlap:] if overlap else ""
                buffer = f"{buffer} {sentence}".strip()
            if buffer:
                current = buffer
            continue

        if current and len(current) + len(paragraph) + 2 > target:
            chunks.append(current)
            tail = current[-overlap:] if overlap else ""
            current = f"{tail}\n\n{paragraph}".strip() if tail else paragraph
        else:
            current = f"{current}\n\n{paragraph}".strip() if current else paragraph

    if current:
        chunks.append(current)

    # Fold a tiny trailing chunk into its predecessor.
    if len(chunks) > 1 and len(chunks[-1]) < MIN_CHUNK_CHARS:
        chunks[-2] = f"{chunks[-2]}\n\n{chunks[-1]}"
        chunks.pop()
    return chunks
